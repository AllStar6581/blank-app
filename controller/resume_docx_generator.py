"""DOCX resume generator — ATS-friendly format for international use."""

import locale
from string import Template

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.run import Run

from controller.resume_controller import (
    Experience,
    ResumePage,
    _extract_skill_name,
)


class InternationalDocxGenerator:
    """Generates a clean, ATS-friendly DOCX resume for international use.

    Args:
        resume: Parsed resume data.
        compact: When *True* (default), produces a 1-2 page document with
            tiered experience display and tighter spacing.  When *False*,
            includes every entry with full detail.
    """

    # -- Compact-mode knobs --------------------------------------------------
    DETAILED_EXPERIENCE_COUNT = 3   # full bullets in compact mode
    COMPACT_MAX_BULLETS = 3         # per detailed entry
    COMPACT_MAX_SKILLS = 8          # per detailed entry
    SKILLS_PER_CATEGORY = 8         # global skills section

    TMPL_LEAD_TO_WEBSITE = Template(
        "There are more items covering $str_years_month_exp of work experience."
        " For details see my site: "
    )

    def __init__(self, resume: ResumePage, *, compact: bool = True):
        self.resume = resume
        self.compact = compact
        self.doc = Document()
        self._setup_document()

    # ------------------------------------------------------------------
    # Document setup
    # ------------------------------------------------------------------

    def _setup_document(self):
        """Set document margins and default font."""
        top_bottom = Cm(1.0) if self.compact else Cm(1.5)
        for section in self.doc.sections:
            section.top_margin = top_bottom
            section.bottom_margin = top_bottom
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

    def generate(self, output_path):
        """Generate and save the DOCX document."""
        locale.setlocale(locale.LC_ALL, "en_US.UTF-8")
        self._add_ats_header()
        self._add_concise_summary()
        self._add_ats_experience()
        self._add_ats_education()
        self._add_ats_skills()
        self._add_ats_languages()
        self.doc.save(output_path)

    # ------------------------------------------------------------------
    # Header
    # ------------------------------------------------------------------

    def _add_ats_header(self):
        """Add name, position, and contact links."""
        name_para = self.doc.add_paragraph()
        name_run = name_para.add_run(self.resume.name.upper())
        name_run.font.size = Pt(14)
        name_run.bold = True

        total_years = self.resume.total_experience_months // 12
        exp_text = f" | {total_years}+ years experience" if total_years > 0 else ""
        position_para = self.doc.add_paragraph()
        position_run = position_para.add_run(
            f"{self.resume.expected_position}{exp_text}"
        )
        position_run.font.size = Pt(11)
        position_run.font.color.rgb = RGBColor(80, 80, 80)

        contact_para = self.doc.add_paragraph()
        if self.resume.tel:
            contact_para.add_run(self.resume.tel)
        if self.resume.email:
            contact_para.add_run(" | ")
            self._create_hyperlink(
                contact_para, self.resume.email, f"mailto:{self.resume.email}"
            )
        if self.resume.website:
            contact_para.add_run(" | ")
            self._create_hyperlink(
                contact_para,
                str(self.resume.website).replace("http://", "").replace("https://", ""),
                str(self.resume.website),
            )

        if self.resume.contacts:
            contact_para.add_run(" | ")
            for i, contact in enumerate(self.resume.contacts[:4]):
                if contact.link:
                    self._create_hyperlink(
                        contact_para, contact.text or contact.name, contact.link
                    )
                else:
                    contact_para.add_run(contact.text or contact.name)
                if i < len(self.resume.contacts[:4]) - 1:
                    contact_para.add_run(" | ")

            for run in contact_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)

        self._add_horizontal_line()

    # ------------------------------------------------------------------
    # Summary
    # ------------------------------------------------------------------

    def _add_concise_summary(self):
        """Add a brief professional summary if available."""
        if not self.resume.about:
            return
        summary_para = self.doc.add_paragraph()
        summary_run = summary_para.add_run(self.resume.about)
        summary_run.font.size = Pt(9)
        summary_run.italic = True
        summary_para.paragraph_format.space_after = Pt(4)

    # ------------------------------------------------------------------
    # Experience
    # ------------------------------------------------------------------

    def _add_ats_experience(self):
        """Add work experience — tiered in compact mode, full in expanded."""
        if not self.resume.exp:
            return

        self._add_section_header("WORK EXPERIENCE")
        all_exp = self.resume.ordered_experience

        if self.compact:
            detailed = all_exp[: self.DETAILED_EXPERIENCE_COUNT]
            remaining = all_exp[self.DETAILED_EXPERIENCE_COUNT :]

            for idx, exp in enumerate(detailed):
                self._render_detailed_entry(exp, max_bullets=self.COMPACT_MAX_BULLETS)

            if remaining:
                self._render_earlier_career(remaining)
        else:
            for exp in all_exp:
                self._render_detailed_entry(exp, max_bullets=None)

    def _render_detailed_entry(self, exp: Experience, *, max_bullets: int | None):
        """Render a single experience entry with configurable bullet limit."""
        # Company | Position | Duration
        header_para = self.doc.add_paragraph()
        header_para.paragraph_format.space_before = Pt(6)
        company_run = header_para.add_run(exp.company_name)
        company_run.bold = True
        company_run.font.size = Pt(10)

        header_para.add_run(" | ")
        header_para.add_run(exp.position_name).font.size = Pt(10)

        header_para.add_run(" | ")
        duration_run = header_para.add_run(self._format_duration(exp))
        duration_run.font.size = Pt(9)
        duration_run.font.color.rgb = RGBColor(120, 120, 120)

        # Description (compact: skip; full: show)
        if not self.compact and exp.description:
            desc_para = self.doc.add_paragraph()
            desc_para.paragraph_format.left_indent = Cm(0.5)
            desc_para.add_run(exp.description).font.size = Pt(9)

        # Merge action_points + responsibilities into one bullet list
        bullets = (exp.action_points or []) + (exp.responsibilities or [])
        for point in bullets[:max_bullets]:
            bp = self.doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Cm(0.5)
            bp.paragraph_format.space_after = Pt(0)
            bp.add_run(point).font.size = Pt(9)

        # Technologies line
        if exp.skills:
            skills_limit = self.COMPACT_MAX_SKILLS if self.compact else None
            names = [_extract_skill_name(s) for s in exp.skills[:skills_limit]]
            skills_para = self.doc.add_paragraph()
            skills_para.paragraph_format.left_indent = Cm(0.5)
            skills_run = skills_para.add_run(f"Technologies: {', '.join(names)}")
            skills_run.font.size = Pt(8)
            skills_run.font.color.rgb = RGBColor(150, 150, 150)
            skills_run.italic = True

    def _render_earlier_career(self, remaining: list[Experience]):
        """Render older positions as compact one-liners under a sub-header."""
        self._add_section_header("EARLIER CAREER")

        for exp in remaining:
            para = self.doc.add_paragraph()
            company_run = para.add_run(exp.company_name)
            company_run.bold = True
            company_run.font.size = Pt(9)

            detail = f" | {exp.position_name} | {self._format_duration(exp)}"
            detail_run = para.add_run(detail)
            detail_run.font.size = Pt(9)
            detail_run.font.color.rgb = RGBColor(100, 100, 100)

        # Link to full details
        if self.resume.website:
            link_para = self.doc.add_paragraph()
            link_para.paragraph_format.space_before = Pt(4)
            run = link_para.add_run("Full details: ")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 100)
            self._create_hyperlink(
                link_para,
                str(self.resume.website).replace("http://", "").replace("https://", ""),
                str(self.resume.website),
            )

    # ------------------------------------------------------------------
    # Education
    # ------------------------------------------------------------------

    def _add_ats_education(self):
        """Add ATS-friendly education section."""
        if not self.resume.edu:
            return

        self._add_section_header("EDUCATION")
        for edu in self.resume.edu[:2]:
            edu_para = self.doc.add_paragraph()
            uni_run = edu_para.add_run(edu.university)
            uni_run.bold = True
            uni_run.font.size = Pt(10)

            edu_para.add_run(" | ")
            edu_para.add_run(edu.degree).font.size = Pt(10)

            edu_para.add_run("\t")
            years_run = edu_para.add_run(f"{edu.year_start} - {edu.year_end}")
            years_run.font.size = Pt(9)
            years_run.font.color.rgb = RGBColor(120, 120, 120)

            if edu.programme:
                prog_para = self.doc.add_paragraph()
                prog_para.paragraph_format.left_indent = Cm(0.5)
                prog_run = prog_para.add_run(edu.programme)
                prog_run.font.size = Pt(9)
                prog_run.font.color.rgb = RGBColor(100, 100, 100)

    # ------------------------------------------------------------------
    # Skills
    # ------------------------------------------------------------------

    def _add_ats_skills(self):
        """Add ATS-friendly skills grouped by category from resume data."""
        categorized = self.resume.all_skills_categorized
        if not categorized:
            return

        self._add_section_header("TECHNICAL SKILLS")
        for category, skills in categorized.items():
            para = self.doc.add_paragraph()
            cat_run = para.add_run(f"{category}: ")
            cat_run.bold = True
            cat_run.font.size = Pt(9)
            para.add_run(
                ", ".join(skills[: self.SKILLS_PER_CATEGORY])
            ).font.size = Pt(9)

    # ------------------------------------------------------------------
    # Languages
    # ------------------------------------------------------------------

    def _add_ats_languages(self):
        """Add spoken languages section."""
        if not self.resume.spoken_languages:
            return
        self._add_section_header("LANGUAGES")
        text = ", ".join(
            f"{lang.name} ({lang.level})" for lang in self.resume.spoken_languages
        )
        self.doc.add_paragraph().add_run(text).font.size = Pt(9)

    # ------------------------------------------------------------------
    # Formatting helpers
    # ------------------------------------------------------------------

    def _format_duration(self, experience: Experience) -> str:
        """Format experience duration as 'Mon YYYY - Mon YYYY'."""
        start = experience.work_start_date_object.strftime("%b %Y")
        end = "Present" if experience.is_still_working else experience.work_end_date_object.strftime("%b %Y")
        return f"{start} - {end}"

    def _add_section_header(self, title: str):
        """Add a bold section header with an underline."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)

        pPr = para._element.get_or_add_pPr()
        border = OxmlElement("w:bottom")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "1")
        border.set(qn("w:color"), "000000")
        pPr.append(border)

    def _add_horizontal_line(self):
        """Add a thin horizontal separator line."""
        para = self.doc.add_paragraph()
        pPr = para._element.get_or_add_pPr()
        border = OxmlElement("w:bottom")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "3")
        border.set(qn("w:space"), "1")
        border.set(qn("w:color"), "000000")
        pPr.append(border)
        para.paragraph_format.space_after = Pt(4)

    def _create_hyperlink(self, paragraph, text: str, url: str) -> Run:
        """Create a clickable blue hyperlink in a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0066CC")
        rPr.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rPr.append(underline)

        new_run.insert(0, rPr)
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

        run_obj = Run(hyperlink.find(qn("w:r")), paragraph)
        run_obj.font.color.rgb = RGBColor(0, 102, 204)
        run_obj.font.underline = True
        run_obj.font.size = Pt(9)
        return run_obj
