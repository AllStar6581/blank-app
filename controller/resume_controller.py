from pydantic import BaseModel, HttpUrl, root_validator
from typing import List, Any
from datetime import date
from dateutil.relativedelta import relativedelta
from matplotlib import font_manager as fm
from dateutil import rrule
import dateparser
from collections import defaultdict
import matplotlib.pyplot as plt
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
import matplotlib.image as mpimg
import qrcode
import vobject
import os
from collections import Counter
from functools import cached_property


class Contact(BaseModel):
    name: str
    icon: str | None
    link: str | None
    text: str | None
    
    def to_markdown(self):
        # return f"[{self.name}]({self.link})"
        return f"[{self.text}]({self.link})"

class Education(BaseModel):
    degree: str
    university: str
    programme: str
    year_start: int
    year_end: int
    website: HttpUrl
    icon: str | None
    
    def to_html(self):
        return ""

class SpokenLanguage(BaseModel):
    name: str
    level: str
    
class Experience(BaseModel):
    
    company_name : str
    company_contacts : str
    company_contacts_link : str | None
    work_start_date : date | str
    work_end_date : date | str | None
    work_start_date_object : date | None
    work_end_date_object : date | None
    total_time_delta : Any
    is_still_working : bool = False
    position_name : str
    description : str | None
    action_points : List[str] | None
    responsibilities : List[str] | None
    skills : List[str] | None
    skill_set_upper : set | None
    video_links: List[str] | None = None
    location: str|None = None

    @property
    def is_still_working(self):
        if not self.work_end_date_object or self.work_end_date_object <= date.today():
            return True
        return False


    @root_validator(pre=True)
    def validate_on_create(cls, values):
        is_still_working = False
        actual_date_start = dateparser.parse(values.get("work_start_date"), date_formats=['%d-%m-%Y', '%m-%Y'], settings={'PREFER_DAY_OF_MONTH': 'first'}).date()
        try:
            actual_date_end = dateparser.parse(values.get("work_end_date"), date_formats=['%d-%m-%Y', '%m-%Y'], settings={'PREFER_DAY_OF_MONTH': 'last'}).date()
        except TypeError:
            actual_date_end = date.today()
            is_still_working = True
        values["work_start_date_object"] = actual_date_start
        values["work_end_date_object"] = actual_date_end
        values["is_still_working"] = is_still_working
        
        delta = relativedelta(actual_date_end, actual_date_start)
        if delta.days >= 0:
            delta.months += 1
        if delta.months >= 12:
            delta.months -= 12
            delta.years += 1
        
        values["total_time_delta"] = delta
        
        
        if is_still_working:
            pass
            # print(values)

        values["skill_set_upper"] = {x.upper() for x in values.get("skills")}
        return values

class ResumePage(BaseModel):
    
    first_name : str = ""
    last_name : str = ""
    email : str | None
    website : HttpUrl | str | None
    expected_position: str = ""
    expected_salary: str | None = None
    photo: str | None  = None
    about: str | None  = None
        
    contacts : List[Contact] | None  = None
    exp : List[Experience]  | None  = None
    edu : List[Education]  | None  = None
    spoken_languages : List[SpokenLanguage] | None = None
    
    @property
    def name(self):
        return f"{self.first_name} {self.last_name}"
    
    @property
    def counter_of_main_skills(self):
        experience_counter = Counter()
        for experience in self.exp:
            for skill in experience.skills:
                experience_counter[skill.lower()] += 1
        return experience_counter
    
    @property
    def ordered_experience(self):
        experience = self.exp
        experience.sort(key=lambda exp_: exp_.work_end_date_object, reverse=True)
        return experience
    
    @property
    def all_skills_set(self):
        experience = self.exp
        skills_set = set()
        for exp_ in experience:
            for skill in exp_.skills:
                skills_set.add(skill)
        return skills_set


    
    # @property
    @cached_property
    def total_experience_months(self):
        year_to_month_work_mapping = defaultdict(set)
        for exp in self.exp:
            date_start = exp.work_start_date_object
            date_end = exp.work_end_date_object
            for dt in rrule.rrule(rrule.MONTHLY, dtstart=date_start, until=date_end):
                year_to_month_work_mapping[dt.year].add(dt.month)

        overall_months = 0
        for year, month_set in year_to_month_work_mapping.items():
            overall_months += len(month_set)
        
        return overall_months
    

    # @property
    @cached_property
    def total_experience_months_wide(self):
        year_to_month_work_mapping = defaultdict(set)
        date_start_list = []
        date_end_list = []
        for exp in self.exp:
            date_start_list.append(exp.work_start_date_object)
            date_end_list.append(exp.work_end_date_object)

        for dt in rrule.rrule(rrule.MONTHLY, dtstart=min(date_start_list), until=max(date_end_list)):
            year_to_month_work_mapping[dt.year].add(dt.month)
        overall_months = 0
        for year, month_set in year_to_month_work_mapping.items():
            overall_months += len(month_set)
        return overall_months
        
    
    @classmethod
    def from_json(cls, json_object):
        contacts_json = json_object.get("contacts")
        contacts = []
        for contact_json in contacts_json:
            contacts.append(Contact(**contact_json))
            
        experience_json = json_object.get("exp")
        exp = []
        for exp_json in experience_json:
            exp.append(Experience(**exp_json))
            
        education_json = json_object.get("edu")
        edu = []
        for edu_json in education_json:
            edu.append(Education(**edu_json))
            
        spoken_languages_json = json_object.get("spoken_languages", [])
        spoken_languages = []
        for spoken_language_json in spoken_languages_json:
            spoken_languages.append(SpokenLanguage(**spoken_language_json))
        
        page_dict = dict(
            contacts=contacts,
            exp=exp,
            edu=edu,
            spoken_languages=spoken_languages,
        )
        page_dict.update(json_object)
        
        return cls(**page_dict)
    
    def to_html(self):
        return ""
    

class PdfPage:
    
    resume_page : ResumePage
    
    # primary_color = "#3E313E"
    # primary_color = "#004d40"
    primary_color = "#00897b"
    # secondary_color = "#007ACC"
    # secondary_color = "#d500f9"
    secondary_color = "#4a148c"
    background_color = "#ffffff"
    black_color = "#000000"
    white_color = "#ffffff"
    

    def __init__(self, resume_page: ResumePage, pallette = {"primary_color" : "#00897b", "secondary_color" : "#4a148c", "background_color" : "#ffffff", "black_color" : "#000000", "white_color" : "#ffffff"}) -> None:
        self.resume_page = resume_page
        if pallette.get("primary_color"):
            self.primary_color = pallette.get("primary_color")
        if pallette.get("secondary_color"):
            self.secondary_color = pallette.get("secondary_color")
        if pallette.get("background_color"):
            self.background_color = pallette.get("background_color")
        if pallette.get("black_color"):
            self.black_color = pallette.get("black_color")
        if pallette.get("white_color"):
            self.white_color = pallette.get("white_color")
        self.qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=4,
        border=0,
    )
    
    def create_vcard_data(self):
        vcard_object = vobject.vCard()
        vcard_object.add('n')
        vcard_object.n.value = vobject.vcard.Name( family=self.resume_page.last_name, given=self.resume_page.first_name )
        vcard_object.add('fn')
        vcard_object.fn.value = self.resume_page.name
        vcard_object.add('email')
        vcard_object.email.value = self.resume_page.email
        vcard_object.email.type_param = 'INTERNET'
        vcard_object.add('title')
        vcard_object.title.value = self.resume_page.expected_position
        # vcard_object.title.type_param = 'TEXT'

        vcard_object.add('url')
        vcard_object.url.value = self.resume_page.website
        vcard_object.url.type_param = 'INTERNET'
        return vcard_object.serialize()
    
    def create_qrcode_image(self, data_string: str, img_output_file_path: str):
        self.qr.add_data(data_string)
        self.qr.make(fit=True)
        # img = self.qr.make_image(fill_color=self.primary_color, back_color=self.background_color)
        img = self.qr.make_image(fill_color=self.primary_color, back_color=self.white_color)
        img.save(img_output_file_path)
        return img
    
    def create_resume_png(self, img_output_file_path: str):
        # plt.rcParams['font.family'] = 'sans-serif'
        # font_path = "fonts/Warownia.otf"
        # font_prop = fm.FontProperties(fname=font_path)
        # plt.rcParams['font.family'] = font_prop.get_name()
        font_path = os.path.join("fonts", "Warownia.otf")
        font_prop = fm.FontProperties(fname=font_path)
        fm.fontManager.addfont(font_path)
        plt.rcParams['font.family'] = font_prop.get_name()
        # plt.rcParams['font.family'] = ["sans-serif"]
        # plt.rcParams['font.sans-serif'] = 'STIXGeneral'
        # plt.rcParams['font.sans-serif'] = 'Helvetica'
        # plt.rcParams['font.sans-serif'] = font_prop.get_name()
        
        # plt.rcParams['axes.facecolor'] = 'none'
        # plt.rcParams['axes.facecolor'] = 'red'
        fig, ax = plt.subplots(figsize=(8.5, 11), facecolor=self.white_color)
        # ax.set_facecolor('white')
        
        # ax.set_facecolor('xkcd:salmon')
        # fig, ax = plt.subplots(figsize=(11.5, 14))

        ax.axvline(x=.5, ymin=0, ymax=1, color=self.secondary_color, alpha=0.0, linewidth=50)
        # plt.axvline(x=.99, color='#000000', alpha=0.5, linewidth=300)
        plt.axvline(x=.99, ymin=0, color=self.primary_color, alpha=1, linewidth=300)
        plt.axhline(y=.835, xmin=0, xmax=1, color=self.white_color, linewidth=3)
        
        exp_position = (.02,.891)
        exp_length = 6
        
        plt.axvline(x=.009, ymin=exp_position[1], ymax=exp_position[1]-exp_length * 0.1 - 0.08, color=self.primary_color, alpha=1, linewidth=2)
        # plt.add_patch(plt.Circle((0.5, 0.5), 0.2, color='blue'))
        # plt.axhline(y=.88, xmin=0, xmax=1, color=background_color, linewidth=3)
        # set background color
        # ax.set_facecolor(self.white_color)
        
        # ax.set_facecolor('red')
        # remove axes
        plt.axis('off')
        # add text
        plt.annotate(self.resume_page.name, (.01,.99), weight='bold', fontsize=20, color=self.primary_color)
        plt.annotate(self.resume_page.expected_position, (.02,.96), weight='regular',color=self.black_color, fontsize=14)
        plt.annotate("\n".join([x.name+":"+x.level for x in self.resume_page.spoken_languages]), (.68,.849), weight='regular', fontsize=8, color=self.background_color)
        
        overall_months = self.resume_page.total_experience_months
        plt.annotate(f"EXPERIENCE | {overall_months // 12} yr.{overall_months % 12 if overall_months % 12 else ''}{' mo' if overall_months % 12 else ''} total", (.02,.920), weight='bold', fontsize=10, color=self.primary_color)

        
        for experience in self.resume_page.ordered_experience[:exp_length]:
            # plt.annotate(experience.company_name, exp_position, weight='bold', fontsize=10, color=self.secondary_color)
            # plt.annotate("●", (-0.001, exp_position[1]), weight='bold', fontsize=10, color=self.primary_color)
            plt.annotate(".", (-0.00, exp_position[1]), weight='bold', fontsize=30, color=self.primary_color)
            plt.annotate(experience.company_name, (exp_position[0]+0.01, exp_position[1]), weight='bold', color=self.black_color, fontsize=10,)
            plt.annotate(experience.position_name, (exp_position[0], exp_position[1] - 0.019), weight='bold', fontsize=10, color=self.secondary_color)
            plt.annotate(f"{experience.work_start_date_object.strftime('%B %Y')} - {experience.work_end_date_object.strftime('%B %Y') if not experience.is_still_working else 'Still working'}  | {str(experience.total_time_delta.years) +' yr' if experience.total_time_delta.years else ''} {str(experience.total_time_delta.months) +' mo' if experience.total_time_delta.months else ''}", (exp_position[0], exp_position[1] - 0.019 - 0.017), weight='regular', fontsize=9, color=self.black_color, alpha=.6)
            description = "\n"
            action_points = experience.action_points[:4]
            for point in action_points:
                description += f"- {point}\n"
            plt.annotate(description, (exp_position[0] + 0.0, exp_position[1] - 0.021 * (len(action_points) + 0.85) - 0.032), weight='regular', color=self.black_color, fontsize=8)
            exp_position = (exp_position[0], exp_position[1] - 0.108 - 0.019)
            exp_position = (exp_position[0], exp_position[1] - 0.019 * (len(experience.action_points) - 4.2)  + -0.016)
        
        plt.annotate("EDUCATION", (.02,.185), weight='bold', fontsize=10, color=self.primary_color)
        
        edu_position = (.02,.155)
        edu_padding = - 0.005
        for education in self.resume_page.edu[:3]:
            plt.annotate(education.university + ", ", edu_position, weight='bold', color=self.black_color, fontsize=10)
            plt.annotate(education.degree, (edu_position[0], edu_position[1] - 0.015 + edu_padding), weight='bold', fontsize=9, color=self.secondary_color)
            plt.annotate(str(education.year_start) + " - " + str(education.year_end), (edu_position[0], edu_position[1] - 0.030 + edu_padding), weight='regular', color=self.black_color, fontsize=9, alpha=.6)
            plt.annotate(education.programme, (edu_position[0], edu_position[1] - 0.045 + edu_padding), weight='regular', color=self.black_color, fontsize=9)
            edu_position = (edu_position[0], edu_position[1] - 0.09 + edu_padding)
        
        plt.annotate("SKILLS", (.7,.80), weight='bold', fontsize=10, color=self.background_color)
        plt.annotate("\n".join([ f"{x[0].capitalize()}" for x in self.resume_page.counter_of_main_skills.most_common(16)]), (.7,.48), weight='regular', fontsize=10, color=self.background_color)
        
        # arr_profile_photo = mpimg.imread('static/me_2.png')
        arr_profile_photo = mpimg.imread('static/me.jpg')
        imagebox_1 = OffsetImage(arr_profile_photo, zoom=0.086)
        ab_1 = AnnotationBbox(imagebox_1, (.90,.92), pad=0.2, bboxprops=dict(edgecolor=self.white_color, facecolor=self.white_color, alpha=1.0))
        ax.add_artist(ab_1)

        plt.annotate(f"ADD MY CONTACT\n& SEE MORE DETAILS:", (.7,.27), weight='bold', fontsize=10, color=self.background_color)

        self.create_qrcode_image(self.create_vcard_data(), "qrcode.png")
        arr_code = mpimg.imread('qrcode.png')
        imagebox = OffsetImage(arr_code, zoom=0.50)
        ab = AnnotationBbox(imagebox, (0.83, 0.15), bboxprops=dict(edgecolor=self.white_color, facecolor=self.white_color, alpha=1.0))
        ax.add_artist(ab)
        # plt.savefig('resumeexample.png', dpi=300, bbox_inches='tight')
        
        plt.annotate(self.resume_page.website or "", (.7,.017), weight='bold', fontsize=10, color=self.background_color)
        
        return plt.savefig(img_output_file_path, dpi=300, bbox_inches='tight') 
 
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

class InternationalDocxGenerator:
    """Generates a clean, professional DOCX resume for international use"""
    
    def __init__(self, resume: ResumePage):
        self.resume = resume
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Setup document styles and margins"""
        # Set narrower margins for better space utilization
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
    
    def generate(self, output_path: str):
        """Generate and save the DOCX document"""
        self._add_header()
        self._add_summary()
        self._add_experience()
        self._add_education()
        self._add_skills()
        self._add_languages()
        
        self.doc.save(output_path)
    
    def _add_header(self):
        """Add name, position, and contacts"""
        # Name
        name_para = self.doc.add_paragraph()
        name_run = name_para.add_run(self.resume.name.upper())
        name_run.font.size = Pt(16)
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Position
        position_para = self.doc.add_paragraph()
        position_run = position_para.add_run(self.resume.expected_position)
        position_run.font.size = Pt(12)
        position_run.font.color.rgb = RGBColor(100, 100, 100)
        position_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contacts
        if self.resume.contacts:
            contacts_para = self.doc.add_paragraph()
            contacts_text = " | ".join(contact.text for contact in self.resume.contacts)
            contacts_run = contacts_para.add_run(contacts_text)
            contacts_run.font.size = Pt(9)
            contacts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Add spacing
    
    def _add_summary(self):
        """Add professional summary"""
        if self.resume.about:
            self._add_section_header("Professional Summary")
            summary_para = self.doc.add_paragraph()
            summary_para.add_run(self.resume.about)
            summary_para.paragraph_format.space_after = Pt(12)
    
    def _add_experience(self):
        """Add work experience section"""
        if not self.resume.exp:
            return
            
        self._add_section_header("Work Experience")
        
        for exp in self.resume.ordered_experience:
            # Company and position line
            exp_para = self.doc.add_paragraph()
            
            # Company (bold)
            company_run = exp_para.add_run(exp.company_name)
            company_run.bold = True
            company_run.font.size = Pt(11)
            
            # Separator
            exp_para.add_run(" | ")
            
            # Position
            position_run = exp_para.add_run(exp.position_name)
            position_run.font.size = Pt(11)
            
            # Duration (right-aligned)
            duration = self._format_duration(exp)
            duration_run = exp_para.add_run(f" | {duration}")
            duration_run.font.size = Pt(10)
            duration_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Location if available
            if exp.location:
                location_run = exp_para.add_run(f" | {exp.location}")
                location_run.font.size = Pt(10)
                location_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Responsibilities
            if exp.responsibilities:
                for responsibility in exp.responsibilities:
                    resp_para = self.doc.add_paragraph(style='List Bullet')
                    resp_para.add_run(responsibility)
                    resp_para.paragraph_format.left_indent = Inches(0.25)
                    resp_para.paragraph_format.space_after = Pt(6)
            
            # Skills used in this position
            if exp.skills:
                skills_para = self.doc.add_paragraph()
                skills_run = skills_para.add_run("Technologies: ")
                skills_run.font.size = Pt(9)
                skills_run.italic = True
                skills_run.font.color.rgb = RGBColor(100, 100, 100)
                
                skills_text = ", ".join(exp.skills)
                skills_value_run = skills_para.add_run(skills_text)
                skills_value_run.font.size = Pt(9)
                skills_value_run.font.color.rgb = RGBColor(100, 100, 100)
            
            self.doc.add_paragraph()  # Add spacing between experiences
    
    def _add_education(self):
        """Add education section"""
        if not self.resume.edu:
            return
            
        self._add_section_header("Education")
        
        for edu in self.resume.edu:
            edu_para = self.doc.add_paragraph()
            
            # University (bold)
            uni_run = edu_para.add_run(edu.university)
            uni_run.bold = True
            uni_run.font.size = Pt(11)
            
            # Degree
            edu_para.add_run(" | ")
            degree_run = edu_para.add_run(edu.degree)
            degree_run.font.size = Pt(11)
            
            # Years
            years_run = edu_para.add_run(f" | {edu.year_start}-{edu.year_end}")
            years_run.font.size = Pt(10)
            years_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Faculty if available
            if edu.programme:
                faculty_para = self.doc.add_paragraph()
                faculty_run = faculty_para.add_run(edu.programme)
                faculty_run.font.size = Pt(10)
                faculty_run.italic = True
                faculty_run.font.color.rgb = RGBColor(100, 100, 100)
                faculty_para.paragraph_format.left_indent = Inches(0.25)
    
    def _add_skills(self):
        """Add skills section with categorization"""
        if not self.resume.all_skills_set:
            return
            
        self._add_section_header("Skills")
        
        # Group skills by category
        # TODO
        # skills_by_category = {}
        # for skill in self.resume.all_skills_set:
        #     if skill.category not in skills_by_category:
        #         skills_by_category[skill.category] = []
        #     skills_by_category[skill.category].append(skill)
        
        # # Create a table for skills
        # if skills_by_category:
        #     table = self.doc.add_table(rows=1, cols=len(skills_by_category))
        #     table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
        #     # Add each category as a column
        #     for i, (category, skills) in enumerate(skills_by_category.items()):
        #         cell = table.cell(0, i)
                
        #         # Category header
        #         category_para = cell.paragraphs[0]
        #         category_run = category_para.add_run(category.title())
        #         category_run.bold = True
        #         category_run.font.size = Pt(10)
                
        #         # Skills list
        #         for skill in skills:
        #             skill_para = cell.add_paragraph()
        #             skill_text = skill.name
        #             if skill.level:
        #                 skill_text += f" ({skill.level})"
        #             if skill.years:
        #                 skill_text += f" - {skill.years}yr"
                    
        #             skill_run = skill_para.add_run(skill_text)
        #             skill_run.font.size = Pt(9)
        #             skill_para.paragraph_format.space_after = Pt(2)
    
    def _add_languages(self):
        """Add languages section"""
        if not self.resume.spoken_languages:
            return
            
        self._add_section_header("Languages")
        
        languages_para = self.doc.add_paragraph()
        languages_text = ", ".join(
            f"{lang.name} ({lang.level})" for lang in self.resume.spoken_languages
        )
        languages_para.add_run(languages_text)
    
    def _add_section_header(self, title: str):
        """Add a section header with formatting"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add underline
        p = para._element
        pPr = p.get_or_add_pPr()
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '8')
        bottom_border.set(qn('w:space'), '1')
        bottom_border.set(qn('w:color'), 'auto')
        pPr.append(bottom_border)
    
    def _format_duration(self, experience: Experience) -> str:
        """Format experience duration"""
        if experience.is_still_working:
            end_text = "Present"
        else:
            end_text = experience.work_end_date_object.strftime("%b %Y")
        
        start_text = experience.work_start_date_object.strftime("%b %Y")
        return f"{start_text} - {end_text}"